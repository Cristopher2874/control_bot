"""
LangChain tools for Word document creation and automation.
Uses python-docx to avoid GUI automation fragility.
"""

import os
from pathlib import Path
from datetime import datetime
from typing import Optional

from langchain_core.tools import tool
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ============================================================================
# TOOL 1: Create a homework report skeleton
# ============================================================================

@tool
def create_homework_skeleton(
    course_name: str,
    assignment_name: str,
    due_date: Optional[str] = None,
    student_name: Optional[str] = None,
    output_dir: Optional[str] = None
) -> str:
    """
    Creates a Word document skeleton for homework assignments.
    
    Args:
        course_name: Name of the course (e.g., "Control Systems")
        assignment_name: Name of the assignment (e.g., "Problem Set 3")
        due_date: Due date (optional, e.g., "Friday, January 17")
        student_name: Your name (optional)
        output_dir: Directory to save the file (default: C:\\Users\\Gokus\\OneDrive\\Escritorio)
    
    Returns:
        Path to the created document and confirmation message.
    
    Example:
        create_homework_skeleton("Control Systems", "Problem Set 3", "Friday")
    """
    try:
        # Default to a standard homework directory
        if output_dir is None:
            output_dir = r"C:\\Users\\Gokus\\OneDrive\\Escritorio"
        
        # Ensure directory exists
        Path(output_dir).mkdir(parents=True, exist_ok=True)
        
        # Create document
        doc = Document()
        
        # Add title section
        title = doc.add_heading(f"{course_name} — {assignment_name}", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add metadata section
        metadata = doc.add_paragraph()
        if student_name:
            metadata.add_run(f"Student: {student_name}\n").bold = True
        if due_date:
            metadata.add_run(f"Due: {due_date}\n").bold = True
        metadata.add_run(f"Date: {datetime.now().strftime('%Y-%m-%d')}\n").bold = True
        
        doc.add_paragraph()  # Spacing
        
        # Add problem sections (placeholder)
        doc.add_heading("Problem 1", level=2)
        doc.add_paragraph("[Insert problem statement here]")
        doc.add_paragraph("[Your solution goes here]")
        doc.add_paragraph()
        
        doc.add_heading("Problem 2", level=2)
        doc.add_paragraph("[Insert problem statement here]")
        doc.add_paragraph("[Your solution goes here]")
        doc.add_paragraph()
        
        doc.add_heading("Problem 3", level=2)
        doc.add_paragraph("[Insert problem statement here]")
        doc.add_paragraph("[Your solution goes here]")
        
        # Save document
        filename = f"{assignment_name.replace(' ', '_')}.docx"
        filepath = os.path.join(output_dir, filename)
        doc.save(filepath)
        
        return f"Created: {filepath}\nReady to edit and submit."
    
    except Exception as e:
        return f"Error creating homework skeleton: {str(e)}"


# ============================================================================
# TOOL 2: Append content to an existing Word document
# ============================================================================

@tool
def append_to_word_document(
    filepath: str,
    section_heading: Optional[str] = None,
    content: Optional[str] = None,
    is_table: bool = False,
    table_rows: Optional[int] = None,
    table_cols: Optional[int] = None
) -> str:
    """
    Appends content to an existing Word document.
    
    Args:
        filepath: Full path to the .docx file
        section_heading: Optional heading for a new section
        content: Text to append
        is_table: If True, create a table instead of text
        table_rows: Number of rows (if is_table=True)
        table_cols: Number of columns (if is_table=True)
    
    Returns:
        Confirmation message.
    
    Example:
        append_to_word_document(
            "C:\\path\\to\\report.docx",
            section_heading="Results",
            content="The experiment showed..."
        )
    """
    try:
        if not os.path.exists(filepath):
            return f"File not found: {filepath}"
        
        doc = Document(filepath)
        
        # Add heading if provided
        if section_heading:
            doc.add_heading(section_heading, level=2)
        
        # Add content or table
        if is_table and table_rows and table_cols:
            table = doc.add_table(rows=table_rows, cols=table_cols)
            table.style = 'Table Grid'
            doc.save(filepath)
            return f"Added {table_rows}x{table_cols} table to {filepath}"
        elif content:
            doc.add_paragraph(content)
            doc.save(filepath)
            return f"Appended content to {filepath}"
        else:
            doc.save(filepath)
            return "No content or table specified."
    
    except Exception as e:
        return f"Error appending to document: {str(e)}"


# ============================================================================
# TOOL 3: Create a lab report template with standard sections
# ============================================================================

@tool
def create_lab_report_template(
    lab_name: str,
    course_name: str,
    output_dir: Optional[str] = None
) -> str:
    """
    Creates a lab report template with standard engineering sections:
    - Objective
    - Procedure
    - Data
    - Analysis
    - Conclusion
    - References
    
    Args:
        lab_name: Name of the lab (e.g., "Resonance in RLC Circuits")
        course_name: Course name (e.g., "Circuit Analysis II")
        output_dir: Directory to save (default: C:\\Users\\you\\School\\Lab_Reports)
    
    Returns:
        Path to created document.
    """
    try:
        if output_dir is None:
            output_dir = r"C:\Users\Cristopher\School\Lab_Reports"
        
        Path(output_dir).mkdir(parents=True, exist_ok=True)
        
        doc = Document()
        
        # Title page style
        title = doc.add_heading(f"Lab Report: {lab_name}", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = doc.add_paragraph(course_name)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_format = subtitle.runs[0]
        subtitle_format.italic = True
        
        doc.add_paragraph(f"Date: {datetime.now().strftime('%B %d, %Y')}")
        doc.add_page_break()
        
        # Standard sections
        sections = [
            ("Objective", "Describe the goal and purpose of this lab."),
            ("Procedure", "Outline the steps taken during the experiment."),
            ("Data", "Present raw data, measurements, and observations."),
            ("Analysis", "Interpret the data and discuss results."),
            ("Conclusion", "Summarize findings and discuss errors/improvements."),
            ("References", "List sources and citations.")
        ]
        
        for heading, placeholder in sections:
            doc.add_heading(heading, level=2)
            doc.add_paragraph(placeholder)
            doc.add_paragraph()  # Spacing
        
        filename = f"{lab_name.replace(' ', '_')}_Report.docx"
        filepath = os.path.join(output_dir, filename)
        doc.save(filepath)
        
        return f"Lab report template created: {filepath}"
    
    except Exception as e:
        return f"Error creating lab report: {str(e)}"


# ============================================================================
# TOOL 4: Open a Word document (for review/editing)
# ============================================================================

@tool
def open_word_document(filepath: str) -> str:
    """
    Opens an existing Word document in your default application.
    
    Args:
        filepath: Full path to the .docx file
    
    Returns:
        Confirmation message.
    """
    try:
        if not os.path.exists(filepath):
            return f"❌ File not found: {filepath}"
        
        # Windows: use 'start' to open with default app
        os.startfile(filepath)
        return f"✅ Opened: {filepath}"
    
    except Exception as e:
        return f"❌ Error opening document: {str(e)}"


# ============================================================================
# Export all tools
# ============================================================================

WORD_TOOLS = [
    create_homework_skeleton,
    append_to_word_document,
    create_lab_report_template,
    open_word_document,
]